#!/usr/bin/env python3
"""Generate supplementary tables/charts used in spreadsheet/doc outputs.

Outputs:
- paper/artifacts/tables/run_manifest_with_timing.csv
- paper/artifacts/tables/latest_version_per_combination.csv
- paper/artifacts/figures/method vs accuracy.png
- paper/artifacts/figures/method vs runtime.png
- output/spreadsheet/all_combinations_table.xlsx
- output/spreadsheet/evaluation_method_comparison_charts.xlsx
- output/spreadsheet/method_summary_for_plots.csv
- output/spreadsheet/key_methods_accuracy.svg
- output/spreadsheet/key_methods_elapsed.svg
- output/spreadsheet/accuracy_vs_time_scatter.svg
- output/doc/best_runs_all_variations_table.docx
"""

from __future__ import annotations

import csv
import json
import statistics
import struct
import zlib
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference, ScatterChart, Series
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Alignment, Font, PatternFill


INDEX_ALIASES = {
    "index_1": "bge_large_en_v1.5",
}


@dataclass
class RunRow:
    run_id: str
    group_key: str
    family: str
    index_name: str
    retrieval_mode: str
    coarse_mode: str
    reranker: str
    reformulation: str
    prompt_mode: str
    llm_model: str
    version: str
    total: int
    correct: int
    accuracy: float
    elapsed_seconds: float
    examples_per_second: float
    seconds_per_example: float
    metrics_path: str
    predictions_path: str

    @property
    def method(self) -> str:
        if self.family == "NO_RAG":
            return f"NO_RAG {self.prompt_mode} ({self.llm_model})"
        return (
            f"RAG {self.retrieval_mode} {self.prompt_mode} "
            f"({self.index_name}, {self.coarse_mode}, {self.reranker}, "
            f"{self.reformulation}, {self.llm_model})"
        )


class Canvas:
    """Tiny no-dependency PNG canvas."""

    def __init__(self, width: int, height: int, bg: Tuple[int, int, int] = (255, 255, 255)):
        self.width = width
        self.height = height
        self.pixels = bytearray(width * height * 3)
        self.clear(bg)

    def clear(self, color: Tuple[int, int, int]) -> None:
        row = bytes(color) * self.width
        for y in range(self.height):
            start = y * self.width * 3
            self.pixels[start : start + self.width * 3] = row

    def set_pixel(self, x: int, y: int, color: Tuple[int, int, int]) -> None:
        if x < 0 or y < 0 or x >= self.width or y >= self.height:
            return
        idx = (y * self.width + x) * 3
        self.pixels[idx : idx + 3] = bytes(color)

    def draw_rect(
        self,
        x0: int,
        y0: int,
        x1: int,
        y1: int,
        color: Tuple[int, int, int],
        fill: bool = True,
    ) -> None:
        x0, x1 = sorted((x0, x1))
        y0, y1 = sorted((y0, y1))
        if fill:
            for y in range(y0, y1 + 1):
                for x in range(x0, x1 + 1):
                    self.set_pixel(x, y, color)
        else:
            for x in range(x0, x1 + 1):
                self.set_pixel(x, y0, color)
                self.set_pixel(x, y1, color)
            for y in range(y0, y1 + 1):
                self.set_pixel(x0, y, color)
                self.set_pixel(x1, y, color)

    def draw_line(self, x0: int, y0: int, x1: int, y1: int, color: Tuple[int, int, int]) -> None:
        dx = abs(x1 - x0)
        sx = 1 if x0 < x1 else -1
        dy = -abs(y1 - y0)
        sy = 1 if y0 < y1 else -1
        err = dx + dy
        while True:
            self.set_pixel(x0, y0, color)
            if x0 == x1 and y0 == y1:
                break
            e2 = 2 * err
            if e2 >= dy:
                err += dy
                x0 += sx
            if e2 <= dx:
                err += dx
                y0 += sy

    def save_png(self, out_path: Path) -> None:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        raw = bytearray()
        stride = self.width * 3
        for y in range(self.height):
            raw.append(0)
            start = y * stride
            raw.extend(self.pixels[start : start + stride])
        compressed = zlib.compress(bytes(raw), level=9)

        def chunk(tag: bytes, data: bytes) -> bytes:
            return (
                struct.pack("!I", len(data))
                + tag
                + data
                + struct.pack("!I", zlib.crc32(tag + data) & 0xFFFFFFFF)
            )

        ihdr = struct.pack("!IIBBBBB", self.width, self.height, 8, 2, 0, 0, 0)
        png = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", compressed) + chunk(b"IEND", b"")
        out_path.write_bytes(png)


def canonical_index(index_name: str) -> str:
    return INDEX_ALIASES.get(index_name, index_name)


def version_rank(version: str) -> Tuple[int, str]:
    v = (version or "base").strip().lower()
    if v == "base":
        return (0, v)
    if v.startswith("v") and v[1:].isdigit():
        return (int(v[1:]), v)
    return (-1, v)


def parse_runs(repo_root: Path) -> List[RunRow]:
    rows: List[RunRow] = []
    final_root = repo_root / "evaluation_results" / "final_output"
    metrics_files = sorted(final_root.glob("**/metrics.json"))

    for metrics_path in metrics_files:
        rel = metrics_path.relative_to(repo_root)
        parts = rel.parts
        with metrics_path.open("r", encoding="utf-8") as f:
            m = json.load(f)

        total = int(m.get("total", 0) or 0)
        correct = int(m.get("correct", 0) or 0)
        accuracy = float(m.get("accuracy", 0.0) or 0.0)
        elapsed = float(m.get("elapsed_seconds", 0.0) or 0.0)
        eps = float(m.get("examples_per_second", 0.0) or 0.0)
        spe = elapsed / total if total > 0 else 0.0
        preds_rel = metrics_path.with_name("predictions.jsonl").relative_to(repo_root)

        if len(parts) < 4 or parts[0] != "evaluation_results" or parts[1] != "final_output":
            continue

        section = parts[2]
        if section == "NO_RAG":
            # NO_RAG/{prompt}/{model}/[vN]/metrics.json
            prompt_mode = parts[3]
            llm_model = parts[4]
            version = "base"
            if len(parts) >= 7 and parts[5].startswith("v"):
                version = parts[5]
            group_key = "|".join(["NO_RAG", prompt_mode, llm_model])
            run_id = f"{group_key}|{version}"
            rows.append(
                RunRow(
                    run_id=run_id,
                    group_key=group_key,
                    family="NO_RAG",
                    index_name="none",
                    retrieval_mode="none",
                    coarse_mode="none",
                    reranker="none",
                    reformulation="none",
                    prompt_mode=prompt_mode,
                    llm_model=llm_model,
                    version=version,
                    total=total,
                    correct=correct,
                    accuracy=accuracy,
                    elapsed_seconds=elapsed,
                    examples_per_second=eps,
                    seconds_per_example=spe,
                    metrics_path=str(rel),
                    predictions_path=str(preds_rel),
                )
            )
        elif section == "RAG" and len(parts) >= 11:
            # RAG/{index}/{retrieval}/{coarse}/{reranker}/{reform}/{model}/{prompt}/[vN]/metrics.json
            index_name = canonical_index(parts[3])
            retrieval_mode = parts[4]
            coarse_mode = parts[5]
            reranker = parts[6]
            reformulation = parts[7]
            llm_model = parts[8]
            prompt_mode = parts[9]
            version = "base"
            if len(parts) >= 12 and parts[10].startswith("v"):
                version = parts[10]
            group_key = "|".join(
                [
                    "RAG",
                    index_name,
                    retrieval_mode,
                    coarse_mode,
                    reranker,
                    reformulation,
                    llm_model,
                    prompt_mode,
                ]
            )
            run_id = f"{group_key}|{version}"
            rows.append(
                RunRow(
                    run_id=run_id,
                    group_key=group_key,
                    family="RAG",
                    index_name=index_name,
                    retrieval_mode=retrieval_mode,
                    coarse_mode=coarse_mode,
                    reranker=reranker,
                    reformulation=reformulation,
                    prompt_mode=prompt_mode,
                    llm_model=llm_model,
                    version=version,
                    total=total,
                    correct=correct,
                    accuracy=accuracy,
                    elapsed_seconds=elapsed,
                    examples_per_second=eps,
                    seconds_per_example=spe,
                    metrics_path=str(rel),
                    predictions_path=str(preds_rel),
                )
            )

    return rows


def write_csv(path: Path, rows: Iterable[Dict[str, object]], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for row in rows:
            w.writerow(row)


def split_latest_and_dropped(rows: List[RunRow]) -> Tuple[List[RunRow], List[RunRow]]:
    grouped: Dict[str, List[RunRow]] = defaultdict(list)
    for r in rows:
        grouped[r.group_key].append(r)

    latest: List[RunRow] = []
    dropped: List[RunRow] = []
    for items in grouped.values():
        selected = max(items, key=lambda x: (version_rank(x.version), x.accuracy, -x.elapsed_seconds))
        latest.append(selected)
        for row in items:
            if row.run_id != selected.run_id:
                dropped.append(row)

    latest.sort(key=lambda x: (x.family, x.index_name, x.prompt_mode, x.llm_model, x.group_key))
    dropped.sort(key=lambda x: (x.group_key, version_rank(x.version)))
    return latest, dropped


def best(rows: List[RunRow], predicate) -> Optional[RunRow]:
    subset = [r for r in rows if predicate(r)]
    if not subset:
        return None
    return max(subset, key=lambda x: x.accuracy)


def build_key_methods(latest: List[RunRow]) -> List[Dict[str, object]]:
    specs = [
        ("NO_RAG zero-shot", lambda r: r.family == "NO_RAG" and r.prompt_mode == "zero_shot"),
        ("NO_RAG CoT", lambda r: r.family == "NO_RAG" and r.prompt_mode == "cot"),
        (
            "RAG dense zs r_off no_ref",
            lambda r: r.family == "RAG"
            and r.retrieval_mode == "dense"
            and r.prompt_mode == "zero_shot"
            and r.reranker == "reranker_off"
            and r.reformulation == "no_reformulation",
        ),
        (
            "RAG dense zs r_on ref",
            lambda r: r.family == "RAG"
            and r.retrieval_mode == "dense"
            and r.prompt_mode == "zero_shot"
            and r.reranker == "reranker_on"
            and r.reformulation == "reformulation",
        ),
        (
            "RAG hybrid zs r_on ref",
            lambda r: r.family == "RAG"
            and r.retrieval_mode == "hybrid"
            and r.prompt_mode == "zero_shot"
            and r.reranker == "reranker_on"
            and r.reformulation == "reformulation",
        ),
        (
            "RAG dense CoT r_on ref",
            lambda r: r.family == "RAG"
            and r.retrieval_mode == "dense"
            and r.prompt_mode == "cot"
            and r.reranker == "reranker_on"
            and r.reformulation == "reformulation",
        ),
    ]

    out: List[Dict[str, object]] = []
    for label, pred in specs:
        item = best(latest, pred)
        if item is None:
            continue
        out.append(
            {
                "method": label,
                "run_id": item.run_id,
                "accuracy": round(item.accuracy, 6),
                "elapsed_seconds": round(item.elapsed_seconds, 1),
                "examples_per_second": round(item.examples_per_second, 3),
            }
        )
    return out


def compute_factor_effects(latest: List[RunRow]) -> List[Dict[str, object]]:
    rag = [r for r in latest if r.family == "RAG"]
    factors = [
        ("reranker", "reranker_off", "reranker_on"),
        ("reformulation", "no_reformulation", "reformulation"),
        ("retrieval_mode", "dense", "hybrid"),
        ("coarse_mode", "coarse_off", "coarse_k20"),
    ]
    dims = [
        "index_name",
        "retrieval_mode",
        "coarse_mode",
        "reranker",
        "reformulation",
        "prompt_mode",
        "llm_model",
    ]

    rows: List[Dict[str, object]] = []
    for factor, level_a, level_b in factors:
        other_dims = [d for d in dims if d != factor]
        grouped: Dict[Tuple[str, ...], Dict[str, RunRow]] = defaultdict(dict)
        for r in rag:
            key = tuple(str(getattr(r, d)) for d in other_dims)
            grouped[key][str(getattr(r, factor))] = r

        acc_deltas: List[float] = []
        elapsed_deltas: List[float] = []
        eps_deltas: List[float] = []
        for vals in grouped.values():
            if level_a in vals and level_b in vals:
                a = vals[level_a]
                b = vals[level_b]
                acc_deltas.append(b.accuracy - a.accuracy)
                elapsed_deltas.append(b.elapsed_seconds - a.elapsed_seconds)
                eps_deltas.append(b.examples_per_second - a.examples_per_second)

        if acc_deltas:
            rows.append(
                {
                    "comparison": f"{level_b} - {level_a}",
                    "n_pairs": len(acc_deltas),
                    "mean_acc_delta": statistics.mean(acc_deltas),
                    "median_acc_delta": statistics.median(acc_deltas),
                    "mean_elapsed_delta_s": statistics.mean(elapsed_deltas),
                    "median_elapsed_delta_s": statistics.median(elapsed_deltas),
                    "mean_eps_delta": statistics.mean(eps_deltas),
                }
            )
        else:
            rows.append(
                {
                    "comparison": f"{level_b} - {level_a}",
                    "n_pairs": 0,
                    "mean_acc_delta": 0.0,
                    "median_acc_delta": 0.0,
                    "mean_elapsed_delta_s": 0.0,
                    "median_elapsed_delta_s": 0.0,
                    "mean_eps_delta": 0.0,
                }
            )
    return rows


def write_svg_bar(
    out_path: Path,
    labels: List[str],
    values: List[float],
    title: str,
    value_fmt: str,
    color: str,
) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1200, 680
    left, right, top, bottom = 90, 1140, 90, 560
    ymax = max(values) if values else 1.0
    if ymax <= 0:
        ymax = 1.0

    bars = []
    n = max(1, len(values))
    for i, v in enumerate(values):
        x0 = left + int(i * (right - left) / n) + 24
        x1 = left + int((i + 1) * (right - left) / n) - 24
        y = bottom - int((v / ymax) * (bottom - top))
        bars.append((x0, y, x1, bottom))

    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">',
        '<rect width="100%" height="100%" fill="white"/>',
        f'<text x="{w//2}" y="38" text-anchor="middle" font-family="Arial" font-size="24">{title}</text>',
        f'<rect x="{left}" y="{top}" width="{right-left}" height="{bottom-top}" fill="none" stroke="#202020" stroke-width="1"/>',
    ]
    for i in range(6):
        gy = top + int(i * (bottom - top) / 5)
        lines.append(
            f'<line x1="{left}" y1="{gy}" x2="{right}" y2="{gy}" stroke="#e5e5e5" stroke-width="1"/>'
        )
    for i, (x0, y, x1, y1) in enumerate(bars):
        v = values[i]
        lines.append(f'<rect x="{x0}" y="{y}" width="{x1-x0}" height="{y1-y}" fill="{color}" opacity="0.85"/>')
        lines.append(f'<text x="{(x0+x1)//2}" y="{y-8}" text-anchor="middle" font-family="Arial" font-size="12">{format(v, value_fmt)}</text>')
        lx = (x0 + x1) // 2
        label = labels[i][:22]
        lines.append(
            f'<text x="{lx}" y="{bottom+18}" transform="rotate(35 {lx},{bottom+18})" '
            'text-anchor="start" font-family="Arial" font-size="11">'
            f"{label}</text>"
        )
    lines.append("</svg>")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def write_svg_scatter(out_path: Path, rows: List[RunRow]) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1200, 700
    left, right, top, bottom = 100, 1140, 80, 610
    points = [(r.elapsed_seconds, r.accuracy, r.family, r.prompt_mode) for r in rows]
    xs = [p[0] for p in points] or [0.0, 1.0]
    ys = [p[1] for p in points] or [0.0, 1.0]
    xmin, xmax = min(xs), max(xs)
    ymin, ymax = min(ys), max(ys)
    if xmax <= xmin:
        xmax = xmin + 1.0
    if ymax <= ymin:
        ymax = ymin + 1.0

    def map_x(v: float) -> int:
        return left + int((v - xmin) / (xmax - xmin) * (right - left))

    def map_y(v: float) -> int:
        return bottom - int((v - ymin) / (ymax - ymin) * (bottom - top))

    colors = {
        ("NO_RAG", "zero_shot"): "#377eb8",
        ("NO_RAG", "cot"): "#4daf4a",
        ("RAG", "zero_shot"): "#e41a1c",
        ("RAG", "cot"): "#984ea3",
    }

    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">',
        '<rect width="100%" height="100%" fill="white"/>',
        '<text x="600" y="36" text-anchor="middle" font-family="Arial" font-size="24">Accuracy vs Time (Latest Versions)</text>',
        f'<rect x="{left}" y="{top}" width="{right-left}" height="{bottom-top}" fill="none" stroke="#202020" stroke-width="1"/>',
    ]
    for i in range(6):
        gy = top + int(i * (bottom - top) / 5)
        lines.append(f'<line x1="{left}" y1="{gy}" x2="{right}" y2="{gy}" stroke="#e5e5e5" stroke-width="1"/>')

    for elapsed, acc, family, prompt in points:
        x = map_x(elapsed)
        y = map_y(acc)
        c = colors.get((family, prompt), "#666666")
        lines.append(f'<circle cx="{x}" cy="{y}" r="5" fill="{c}" fill-opacity="0.85"/>')

    lines.append(f'<text x="{(left+right)//2}" y="{h-22}" text-anchor="middle" font-family="Arial" font-size="14">Elapsed seconds</text>')
    lines.append(
        f'<text x="28" y="{(top+bottom)//2}" transform="rotate(-90 28,{(top+bottom)//2})" text-anchor="middle" font-family="Arial" font-size="14">Accuracy</text>'
    )
    lines.append("</svg>")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def write_png_bars(path: Path, values: List[float], mode: str) -> None:
    c = Canvas(960, 560, (255, 255, 255))
    left, right, top, bottom = 80, 920, 60, 500
    axis = (20, 20, 20)
    c.draw_rect(left, top, right, bottom, axis, fill=False)
    vmax = max(values) if values else 1.0
    if vmax <= 0:
        vmax = 1.0
    color = (62, 114, 196) if mode == "accuracy" else (214, 125, 61)
    for i, v in enumerate(values):
        x0 = left + int(i * (right - left) / max(1, len(values))) + 18
        x1 = left + int((i + 1) * (right - left) / max(1, len(values))) - 18
        y = bottom - int((v / vmax) * (bottom - top))
        c.draw_rect(x0, y, x1, bottom, color, fill=True)
        c.draw_rect(x0, y, x1, bottom, axis, fill=False)
    c.save_png(path)


def set_header_style(ws, header_row: int = 1) -> None:
    fill = PatternFill("solid", fgColor="E6E6E6")
    for cell in ws[header_row]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.fill = fill


def autofit_columns(ws, max_width: int = 44) -> None:
    for col in ws.columns:
        letter = col[0].column_letter
        width = max(len(str(c.value)) if c.value is not None else 0 for c in col) + 2
        ws.column_dimensions[letter].width = min(max_width, max(10, width))


def write_all_combinations_workbook(out_path: Path, latest: List[RunRow], dropped: List[RunRow]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "All_Combinations_Latest"
    headers = [
        "method",
        "run_id",
        "family",
        "index_name",
        "retrieval_mode",
        "coarse_mode",
        "reranker",
        "reformulation",
        "prompt_mode",
        "llm_model",
        "version",
        "accuracy",
        "elapsed_seconds",
        "examples_per_second",
        "seconds_per_example",
        "correct",
        "total",
    ]
    ws.append(headers)
    for r in latest:
        ws.append(
            [
                r.method,
                r.run_id,
                r.family,
                r.index_name,
                r.retrieval_mode,
                r.coarse_mode,
                r.reranker,
                r.reformulation,
                r.prompt_mode,
                r.llm_model,
                r.version,
                r.accuracy,
                r.elapsed_seconds,
                r.examples_per_second,
                r.seconds_per_example,
                r.correct,
                r.total,
            ]
        )
    set_header_style(ws)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:Q{ws.max_row}"
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        row[11].number_format = "0.000000"
        row[12].number_format = "0.0"
        row[13].number_format = "0.000"
        row[14].number_format = "0.000"
    ws.conditional_formatting.add(
        f"L2:L{ws.max_row}",
        ColorScaleRule(
            start_type="min",
            start_color="F4C7C3",
            mid_type="percentile",
            mid_value=50,
            mid_color="FFEB84",
            end_type="max",
            end_color="63BE7B",
        ),
    )
    autofit_columns(ws)

    ws_drop = wb.create_sheet("Dropped_Older_Versions")
    ws_drop.append(headers)
    for r in dropped:
        ws_drop.append(
            [
                r.method,
                r.run_id,
                r.family,
                r.index_name,
                r.retrieval_mode,
                r.coarse_mode,
                r.reranker,
                r.reformulation,
                r.prompt_mode,
                r.llm_model,
                r.version,
                r.accuracy,
                r.elapsed_seconds,
                r.examples_per_second,
                r.seconds_per_example,
                r.correct,
                r.total,
            ]
        )
    set_header_style(ws_drop)
    ws_drop.freeze_panes = "A2"
    ws_drop.auto_filter.ref = f"A1:Q{max(2, ws_drop.max_row)}"
    autofit_columns(ws_drop)

    ws_info = wb.create_sheet("README")
    ws_info["A1"] = "All combinations table (latest versions only)"
    ws_info["A3"] = (
        f"All_Combinations_Latest: {len(latest)} rows (one per configuration, latest version kept)"
    )
    ws_info["A4"] = f"Dropped_Older_Versions: {len(dropped)} rows (older reruns removed)"
    ws_info["A6"] = "Version precedence: v4 > v3 > v2 > base (numeric vN preferred over base)."
    ws_info["A1"].font = Font(bold=True, size=13)
    ws_info.column_dimensions["A"].width = 120

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def write_eval_charts_workbook(
    out_path: Path,
    latest: List[RunRow],
    method_rows: List[Dict[str, object]],
    factor_rows: List[Dict[str, object]],
) -> None:
    wb = Workbook()
    ws_best = wb.active
    ws_best.title = "BestRuns"
    ws_best.append(
        [
            "run_id",
            "family",
            "index_name",
            "retrieval_mode",
            "coarse_mode",
            "reranker",
            "reformulation",
            "prompt_mode",
            "llm_model",
            "accuracy",
            "elapsed_seconds",
            "examples_per_second",
        ]
    )
    for r in sorted(latest, key=lambda x: x.accuracy, reverse=True):
        ws_best.append(
            [
                r.run_id,
                r.family,
                r.index_name,
                r.retrieval_mode,
                r.coarse_mode,
                r.reranker,
                r.reformulation,
                r.prompt_mode,
                r.llm_model,
                r.accuracy,
                r.elapsed_seconds,
                r.examples_per_second,
            ]
        )
    set_header_style(ws_best)
    ws_best.freeze_panes = "A2"
    ws_best.auto_filter.ref = f"A1:L{ws_best.max_row}"
    autofit_columns(ws_best)

    ws_key = wb.create_sheet("KeyMethods")
    ws_key.append(["method", "run_id", "accuracy", "elapsed_seconds", "examples_per_second"])
    for row in method_rows:
        ws_key.append(
            [
                row["method"],
                row["run_id"],
                row["accuracy"],
                row["elapsed_seconds"],
                row["examples_per_second"],
            ]
        )
    set_header_style(ws_key)
    autofit_columns(ws_key)

    ws_factor = wb.create_sheet("FactorEffects")
    ws_factor.append(
        [
            "comparison",
            "n_pairs",
            "mean_acc_delta",
            "median_acc_delta",
            "mean_elapsed_delta_s",
            "median_elapsed_delta_s",
            "mean_eps_delta",
        ]
    )
    for row in factor_rows:
        ws_factor.append(
            [
                row["comparison"],
                row["n_pairs"],
                row["mean_acc_delta"],
                row["median_acc_delta"],
                row["mean_elapsed_delta_s"],
                row["median_elapsed_delta_s"],
                row["mean_eps_delta"],
            ]
        )
    set_header_style(ws_factor)
    autofit_columns(ws_factor)

    ws_scatter = wb.create_sheet("ScatterData")
    ws_scatter.append(
        [
            "family",
            "prompt_mode",
            "run_id",
            "elapsed_seconds",
            "accuracy",
            None,
            "NO_RAG zero-shot elapsed",
            "NO_RAG zero-shot acc",
            "NO_RAG CoT elapsed",
            "NO_RAG CoT acc",
            "RAG zero-shot elapsed",
            "RAG zero-shot acc",
            "RAG CoT elapsed",
            "RAG CoT acc",
        ]
    )
    sorted_rows = sorted(latest, key=lambda x: x.accuracy, reverse=True)
    no_rag_zero = [r for r in sorted_rows if r.family == "NO_RAG" and r.prompt_mode == "zero_shot"]
    no_rag_cot = [r for r in sorted_rows if r.family == "NO_RAG" and r.prompt_mode == "cot"]
    rag_zero = [r for r in sorted_rows if r.family == "RAG" and r.prompt_mode == "zero_shot"]
    rag_cot = [r for r in sorted_rows if r.family == "RAG" and r.prompt_mode == "cot"]

    n = max(len(sorted_rows), len(no_rag_zero), len(no_rag_cot), len(rag_zero), len(rag_cot))
    for i in range(n):
        base = sorted_rows[i] if i < len(sorted_rows) else None
        nz = no_rag_zero[i] if i < len(no_rag_zero) else None
        nc = no_rag_cot[i] if i < len(no_rag_cot) else None
        rz = rag_zero[i] if i < len(rag_zero) else None
        rc = rag_cot[i] if i < len(rag_cot) else None
        ws_scatter.append(
            [
                base.family if base else None,
                base.prompt_mode if base else None,
                base.run_id if base else None,
                base.elapsed_seconds if base else None,
                base.accuracy if base else None,
                None,
                nz.elapsed_seconds if nz else None,
                nz.accuracy if nz else None,
                nc.elapsed_seconds if nc else None,
                nc.accuracy if nc else None,
                rz.elapsed_seconds if rz else None,
                rz.accuracy if rz else None,
                rc.elapsed_seconds if rc else None,
                rc.accuracy if rc else None,
            ]
        )
    set_header_style(ws_scatter)
    autofit_columns(ws_scatter)

    ws_charts = wb.create_sheet("Charts")
    ws_charts["A1"] = "Evaluation Method Comparison Charts"
    ws_charts["A2"] = "Dataset size: 1,273 complete examples per run"
    ws_charts["A1"].font = Font(bold=True, size=14)

    if ws_key.max_row >= 2:
        acc_chart = BarChart()
        acc_chart.title = "Method vs Accuracy"
        acc_chart.y_axis.title = "Accuracy"
        acc_chart.x_axis.title = "Method"
        data = Reference(ws_key, min_col=3, min_row=1, max_row=ws_key.max_row)
        cats = Reference(ws_key, min_col=1, min_row=2, max_row=ws_key.max_row)
        acc_chart.add_data(data, titles_from_data=True)
        acc_chart.set_categories(cats)
        acc_chart.height = 7
        acc_chart.width = 12
        ws_charts.add_chart(acc_chart, "A4")

        time_chart = BarChart()
        time_chart.title = "Method vs Runtime (seconds)"
        time_chart.y_axis.title = "Elapsed seconds"
        time_chart.x_axis.title = "Method"
        data2 = Reference(ws_key, min_col=4, min_row=1, max_row=ws_key.max_row)
        time_chart.add_data(data2, titles_from_data=True)
        time_chart.set_categories(cats)
        time_chart.height = 7
        time_chart.width = 12
        ws_charts.add_chart(time_chart, "A20")

    scatter_chart = ScatterChart()
    scatter_chart.title = "Accuracy vs Time"
    scatter_chart.x_axis.title = "Elapsed seconds"
    scatter_chart.y_axis.title = "Accuracy"
    scatter_chart.height = 7
    scatter_chart.width = 12
    max_row = ws_scatter.max_row
    for cx, cy, title in [
        (7, 8, "NO_RAG zero-shot"),
        (9, 10, "NO_RAG CoT"),
        (11, 12, "RAG zero-shot"),
        (13, 14, "RAG CoT"),
    ]:
        xvals = Reference(ws_scatter, min_col=cx, min_row=2, max_row=max_row)
        yvals = Reference(ws_scatter, min_col=cy, min_row=2, max_row=max_row)
        series = Series(yvals, xvals, title=title)
        scatter_chart.series.append(series)
    ws_charts.add_chart(scatter_chart, "N4")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def write_doc_table(out_path: Path, latest: List[RunRow]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    heading = doc.add_paragraph("BestRuns - All Variations")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in heading.runs:
        r.bold = True
        r.font.size = Pt(14)

    columns = [
        "Family",
        "Index",
        "Retrieval",
        "Coarse",
        "Reranker",
        "Reformulation",
        "Prompt",
        "Model",
        "Version",
        "Accuracy",
        "Elapsed (s)",
        "Ex/s",
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.size = Pt(10)
        shade_cell(hdr[i], "E6E6E6")

    ordered = sorted(latest, key=lambda x: (x.family, -x.accuracy, x.index_name, x.prompt_mode))
    for row in ordered:
        cells = table.add_row().cells
        values = [
            row.family,
            row.index_name,
            row.retrieval_mode,
            row.coarse_mode,
            row.reranker,
            row.reformulation,
            row.prompt_mode,
            row.llm_model,
            row.version,
            f"{row.accuracy:.4f}",
            f"{row.elapsed_seconds:.1f}",
            f"{row.examples_per_second:.3f}",
        ]
        for i, val in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    doc.add_paragraph(
        "All rows from latest-version combinations are included. "
        "Metrics are recomputed from current evaluation results."
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    repo_root = Path(__file__).resolve().parents[3]
    table_dir = repo_root / "paper" / "artifacts" / "tables"
    fig_dir = repo_root / "paper" / "artifacts" / "figures"
    sheet_dir = repo_root / "output" / "spreadsheet"
    doc_dir = repo_root / "output" / "doc"

    runs = parse_runs(repo_root)
    if not runs:
        raise RuntimeError("No runs found under evaluation_results/final_output")
    runs = [r for r in runs if r.total == 1273]
    if not runs:
        raise RuntimeError("No completed runs found (total==1273)")

    latest, dropped = split_latest_and_dropped(runs)

    run_manifest_rows = [
        {
            "run_id": r.run_id,
            "family": r.family,
            "index_name": r.index_name,
            "embedding_model": (
                "BAAI/bge-large-en-v1.5"
                if r.index_name == "bge_large_en_v1.5"
                else ("abhinand/MedEmbed-large-v0.1" if r.index_name == "medembed" else "none")
            ),
            "retrieval_mode": r.retrieval_mode,
            "coarse_mode": r.coarse_mode,
            "reranker": r.reranker,
            "reformulation": r.reformulation,
            "prompt_mode": r.prompt_mode,
            "llm_model": r.llm_model,
            "total": r.total,
            "correct": r.correct,
            "accuracy": f"{r.accuracy:.12f}",
            "metrics_path": r.metrics_path,
            "predictions_path": r.predictions_path,
            "is_complete": True,
            "version": r.version,
            "group_key": r.group_key,
            "elapsed_seconds": r.elapsed_seconds,
            "examples_per_second": r.examples_per_second,
            "seconds_per_example": r.seconds_per_example,
        }
        for r in sorted(runs, key=lambda x: (x.family, x.group_key, version_rank(x.version)))
    ]
    write_csv(
        table_dir / "run_manifest_with_timing.csv",
        run_manifest_rows,
        [
            "run_id",
            "family",
            "index_name",
            "embedding_model",
            "retrieval_mode",
            "coarse_mode",
            "reranker",
            "reformulation",
            "prompt_mode",
            "llm_model",
            "total",
            "correct",
            "accuracy",
            "metrics_path",
            "predictions_path",
            "is_complete",
            "version",
            "group_key",
            "elapsed_seconds",
            "examples_per_second",
            "seconds_per_example",
        ],
    )

    latest_rows_for_csv = sorted(latest, key=lambda x: x.accuracy, reverse=True)
    write_csv(
        table_dir / "latest_version_per_combination.csv",
        [
            {
                "family": r.family,
                "index_name": r.index_name,
                "retrieval_mode": r.retrieval_mode,
                "coarse_mode": r.coarse_mode,
                "reranker": r.reranker,
                "reformulation": r.reformulation,
                "prompt_mode": r.prompt_mode,
                "llm_model": r.llm_model,
                "kept_version": r.version,
                "accuracy": f"{r.accuracy:.12f}",
                "correct": r.correct,
                "total": r.total,
                "metrics_path": r.metrics_path,
            }
            for r in latest_rows_for_csv
        ],
        [
            "family",
            "index_name",
            "retrieval_mode",
            "coarse_mode",
            "reranker",
            "reformulation",
            "prompt_mode",
            "llm_model",
            "kept_version",
            "accuracy",
            "correct",
            "total",
            "metrics_path",
        ],
    )

    method_rows = build_key_methods(latest_rows_for_csv)
    write_csv(
        sheet_dir / "method_summary_for_plots.csv",
        method_rows,
        ["method", "run_id", "accuracy", "elapsed_seconds", "examples_per_second"],
    )

    labels = [str(r["method"]) for r in method_rows]
    acc_values = [float(r["accuracy"]) for r in method_rows]
    time_values = [float(r["elapsed_seconds"]) for r in method_rows]
    write_svg_bar(
        sheet_dir / "key_methods_accuracy.svg",
        labels,
        acc_values,
        "Key Methods: Accuracy",
        ".4f",
        "#3e72c4",
    )
    write_svg_bar(
        sheet_dir / "key_methods_elapsed.svg",
        labels,
        time_values,
        "Key Methods: Elapsed Seconds",
        ".1f",
        "#d67d3d",
    )
    write_svg_scatter(sheet_dir / "accuracy_vs_time_scatter.svg", latest_rows_for_csv)

    write_png_bars(fig_dir / "method vs accuracy.png", acc_values, mode="accuracy")
    write_png_bars(fig_dir / "method vs runtime.png", time_values, mode="runtime")

    write_all_combinations_workbook(sheet_dir / "all_combinations_table.xlsx", latest_rows_for_csv, dropped)
    factor_rows = compute_factor_effects(latest_rows_for_csv)
    write_eval_charts_workbook(
        sheet_dir / "evaluation_method_comparison_charts.xlsx",
        latest_rows_for_csv,
        method_rows,
        factor_rows,
    )
    write_doc_table(doc_dir / "best_runs_all_variations_table.docx", latest_rows_for_csv)

    print("Generated supplementary artifacts:")
    print(f"  Tables: {table_dir / 'run_manifest_with_timing.csv'}")
    print(f"  Tables: {table_dir / 'latest_version_per_combination.csv'}")
    print(f"  Spreadsheet: {sheet_dir}")
    print(f"  Doc: {doc_dir / 'best_runs_all_variations_table.docx'}")
    print(f"  Figures: {fig_dir / 'method vs accuracy.png'}")
    print(f"  Figures: {fig_dir / 'method vs runtime.png'}")


if __name__ == "__main__":
    main()
